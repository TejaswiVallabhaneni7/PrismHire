# utils/parser.py
import os
from io import BytesIO
from PyPDF2 import PdfReader
import docx

def extract_text_from_pdf(path_or_stream):
    """Return text from a PDF file path or file-like object."""
    try:
        if hasattr(path_or_stream, "read"):
            reader = PdfReader(path_or_stream)
        else:
            reader = PdfReader(open(path_or_stream, "rb"))
        text = []
        for page in reader.pages:
            txt = page.extract_text()
            if txt:
                text.append(txt)
        return "\n".join(text)
    except Exception as e:
        return ""

def extract_text_from_docx(path_or_stream):
    """Return text from a DOCX file path or file-like object."""
    try:
        if hasattr(path_or_stream, "read"):
            # docx needs a path, so save to a temp BytesIO first
            temp = BytesIO(path_or_stream.read())
            document = docx.Document(temp)
        else:
            document = docx.Document(path_or_stream)
        fullText = []
        for para in document.paragraphs:
            fullText.append(para.text)
        return "\n".join(fullText)
    except Exception as e:
        return ""

def extract_text_from_txt(path_or_stream):
    try:
        if hasattr(path_or_stream, "read"):
            return path_or_stream.read().decode(errors='ignore')
        else:
            with open(path_or_stream, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
    except Exception:
        return ""

def extract_text(filepath):
    """Detect file type by extension and extract text. filepath can be path or file-like object with filename attribute."""
    # If it's a file-like (werkzeug FileStorage), it has .filename and supports read/seek
    if hasattr(filepath, "filename"):
        filename = filepath.filename
        filepath.stream.seek(0)
        ext = filename.lower().split('.')[-1]
        stream = filepath.stream
        if ext == 'pdf':
            text = extract_text_from_pdf(stream)
            stream.seek(0)
            return text
        elif ext in ('docx', 'doc'):
            text = extract_text_from_docx(stream)
            stream.seek(0)
            return text
        elif ext in ('txt', 'text'):
            text = extract_text_from_txt(stream)
            stream.seek(0)
            return text
    else:
        # assume it's a path string
        ext = filepath.lower().split('.')[-1]
        if ext == 'pdf':
            return extract_text_from_pdf(filepath)
        elif ext in ('docx', 'doc'):
            return extract_text_from_docx(filepath)
        elif ext in ('txt', 'text'):
            return extract_text_from_txt(filepath)
    return ""
